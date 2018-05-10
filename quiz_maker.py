# -*- coding: UTF-8 -*-
from __future__ import unicode_literals

import sys

from docx.enum.text import WD_BREAK_TYPE
from docx.oxml import CT_Num, CT_Numbering
from docx.parts.numbering import _NumberingDefinitions

reload(sys)
sys.setdefaultencoding('utf-8')
import csv
from collections import OrderedDict

from docx import Document
from docx.opc.oxml import qn
from docx.shared import Inches, Pt


class Question:
    def __init__(self, question, answers=None, choices=None):
        self.question = question
        self.choices = choices
        self.answers = self._load_config_answers(answers)

    def answer(self, q_answer):
        self.q_answer = self._get_answers(q_answer)

    def eval_answer(self):
        r = self._do_eval_answer()
        if r == 0:
            print "answer: %s" % self.q_answer
            print "correct answer: %s" % self.answers
        return r

    def _do_eval_answer(self):
        if self.q_answer == self.answers:
            return 1
        return 0

    def _load_config_answers(self, answers):
        return answers

    def _get_answers(self, q_answer):
        return q_answer


class ChoiceQuestion(Question):
    def _do_eval_answer(self):
        if len(self.q_answer) == len(self.answers):
            for ans in self.q_answer:
                if ans not in self.answers:
                    return 0
            return 1
        else:
            return 0

    def _load_config_answers(self, answers):
        config_answers = answers.split("|")
        correct_answers = []
        for answer in config_answers:
            correct_answers.append(chr(65 + int(answer)))
        return correct_answers

    def _get_answers(self, q_answer):
        qu_answer = []
        for i in q_answer:
            qu_answer.append(i)
        return qu_answer


def load(path):
    questions = []
    with open(path, 'rb') as m_file:
        reader = csv.reader(m_file)
        contents = reader
        for row in contents:
            question = ""
            choices = []
            answers = ""
            for i in range(0, len(row)):
                if i == 0:
                    question = row[i]
                if i > 0 and "ANS|" not in row[i]:
                    choices.append(row[i])
                if "ANS|" in row[i]:
                    answers = row[i].split("ANS|")[1]
            if len(choices) > 1:
                questions.append(ChoiceQuestion(question, answers, choices))
            else:
                questions.append(Question(question, answers, choices))
    return questions


def test(questions):
    score = 0
    for q in questions:
        print q.question
        choice = 65
        for c in q.choices:
            print chr(choice) + ") " + c
            choice += 1
        q.answer(raw_input("> "))
        if q.eval_answer() == 1:
            score += 1
        print ""
    print "Score: " + str(score) + "/" + str(len(questions))


def paper(questions, sort=True, paper_path="quiz.doc"):
    document = Document()
    for s in document.styles:
        if hasattr(s, 'font'):
            s.font.name = '宋体'
        # s._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    answer_path = "%s_answer.txt" % paper_path
    i = 1
    with open(answer_path, 'w') as f:
        if sort:
            sort_ret = _get_sorted_questions(questions)
            for item in sort_ret.keys():
                if len(sort_ret[item]) > 0:
                    f.write("%s\n" % item)
                    document.add_heading(item, 2)
                    i = _print_questions(sort_ret[item], document, f, i, 2)
                    f.write("\n")
        else:
            _print_questions(questions, document, f)
    document.save(paper_path)


def _get_sorted_questions(questions):
    single_choices_questions = []
    multiple_choices_questions = []
    judge_questions = []
    fill_questions = []
    ret = OrderedDict()
    ret['单选题'] = single_choices_questions
    ret['多选题'] = multiple_choices_questions
    ret['判断题'] = judge_questions
    ret['填空题'] = fill_questions
    for q in questions:
        if len(q.choices) > 1:
            if len(q.answers) > 1:
                multiple_choices_questions.append(q)
            else:
                single_choices_questions.append(q)
        elif q.answers == "Y" or q.answers == "N":
            judge_questions.append(q)
        else:
            fill_questions.append(q)
    return ret


def _print_questions(questions, document, f, start=1, level=1):
    i = start
    questions_style = "List Number"
    num_id = None
    if level > 1:
        questions_style = "%s %d" % (questions_style, level)
        # # Restart numbering of an ordered list
        # next_num_id = document.part.numbering_part.numbering_definitions._numbering._next_numId
        # num = CT_Num.new(1, str(level))
        # num.add_lvlOverride(ilvl=0).add_startOverride(1)
        # num_id = document.part.numbering_part.numbering_definitions._numbering._insert_num(num)
    choices_style = "List %d" % (level + 1)
    for q in questions:
        question_p = document.add_paragraph(unicode(q.question, 'utf-8'), style=questions_style)
        if len(q.choices) > 1:
            choice = 65
            for c in q.choices:
                p = document.add_paragraph(chr(choice) + ") " + unicode(c, 'utf-8'), style=choices_style)
                paragraph_format = p.paragraph_format
                paragraph_format.line_spacing = Pt(20)
                # paragraph_format.left_indent = Inches(0.3)
                choice += 1
        else:
            document.add_paragraph(style=choices_style)
        # question_p.paragraph_format.line_spacing = Pt(20)
        f.write("%d. %s\n" % (i, q.answers))
        i += 1
    return i

questions = load("q.csv")
print "Questions loaded: " + str(len(questions))

paper(questions)
